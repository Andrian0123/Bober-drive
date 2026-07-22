#!/usr/bin/env python3
"""
Graphify Engine - Week 5 Implementation
Document ingestion and knowledge graph enrichment
Parses PDF, DOCX, TXT, Markdown files and creates semantic graph entries

Architecture:
- Multi-format document parsing (PDF/DOCX/TXT/Markdown)
- Section extraction and hierarchical organization
- Automatic entity/keyword extraction
- Integration with VaultCore for storage
- FTS5 indexing for full-text search
- Semantic relationships between documents
"""

import logging
import json
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict, field
from datetime import datetime
from enum import Enum
import re

logger = logging.getLogger(__name__)


class DocumentFormat(Enum):
    """Supported document formats"""
    PDF = "pdf"
    DOCX = "docx"
    MARKDOWN = "markdown"
    TEXT = "text"
    HTML = "html"


class SectionType(Enum):
    """Types of document sections"""
    TITLE = "title"
    HEADING = "heading"
    SUBHEADING = "subheading"
    PARAGRAPH = "paragraph"
    CODE_BLOCK = "code_block"
    LIST = "list"
    TABLE = "table"
    QUOTE = "quote"
    METADATA = "metadata"


@dataclass
class DocumentSection:
    """Represents a section within a document"""
    section_id: str
    section_type: SectionType
    title: str
    content: str
    level: int  # Heading level (1-6)
    parent_section_id: Optional[str] = None
    keywords: List[str] = field(default_factory=list)
    entities: Dict[str, List[str]] = field(default_factory=dict)  # type -> [entities]
    line_range: Tuple[int, int] = field(default_factory=lambda: (0, 0))
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            **asdict(self),
            "section_type": self.section_type.value
        }


@dataclass
class ParsedDocument:
    """Represents a fully parsed document"""
    document_id: str
    format: DocumentFormat
    title: str
    author: Optional[str] = None
    source_path: Optional[str] = None
    content: str = ""
    sections: List[DocumentSection] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    keywords: List[str] = field(default_factory=list)
    entities: Dict[str, List[str]] = field(default_factory=dict)
    created_at: str = field(default_factory=lambda: datetime.utcnow().isoformat())
    parsed_at: str = field(default_factory=lambda: datetime.utcnow().isoformat())
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            **asdict(self),
            "format": self.format.value,
            "sections": [s.to_dict() for s in self.sections]
        }


class GraphifyEngine:
    """Main document ingestion and graph building engine"""
    
    # Entity extraction patterns
    ENTITY_PATTERNS = {
        "email": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        "url": r'https?://[^\s]+',
        "code": r'`[^`]+`',
        "filename": r'(?:^|[\s])([a-zA-Z0-9._/-]+\.[a-zA-Z]{2,4})',
        "function": r'def\s+([a-zA-Z_][a-zA-Z0-9_]*)|function\s+([a-zA-Z_][a-zA-Z0-9_]*)',
        "class": r'class\s+([a-zA-Z_][a-zA-Z0-9_]*)',
    }
    
    def __init__(self, vault_core=None, fts5_extension=None):
        """
        Initialize Graphify Engine
        
        Args:
            vault_core: VaultCore instance for storage
            fts5_extension: FTS5Extension instance for indexing
        """
        self.vault_core = vault_core
        self.fts5_extension = fts5_extension
        self.parsed_documents: Dict[str, ParsedDocument] = {}
        logger.info("GraphifyEngine initialized")
    
    def import_document(self, file_path: Path) -> Optional[ParsedDocument]:
        """
        Import and parse a document
        
        Args:
            file_path: Path to document file
            
        Returns:
            ParsedDocument if successful, None otherwise
        """
        try:
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                return None
            
            # Determine format
            suffix = file_path.suffix.lower()
            if suffix == '.pdf':
                fmt = DocumentFormat.PDF
            elif suffix == '.docx':
                fmt = DocumentFormat.DOCX
            elif suffix in ['.md', '.markdown']:
                fmt = DocumentFormat.MARKDOWN
            elif suffix == '.html':
                fmt = DocumentFormat.HTML
            else:
                fmt = DocumentFormat.TEXT
            
            logger.info(f"Importing {fmt.value} document: {file_path}")
            
            # Parse based on format
            if fmt == DocumentFormat.PDF:
                parsed = self._parse_pdf(file_path)
            elif fmt == DocumentFormat.DOCX:
                parsed = self._parse_docx(file_path)
            elif fmt == DocumentFormat.MARKDOWN:
                parsed = self._parse_markdown(file_path)
            elif fmt == DocumentFormat.HTML:
                parsed = self._parse_html(file_path)
            else:
                parsed = self._parse_text(file_path)
            
            if parsed:
                self.parsed_documents[parsed.document_id] = parsed
                
                # Extract entities and keywords
                self._extract_entities(parsed)
                
                # Save to VaultCore if available
                if self.vault_core:
                    self._save_to_vault(parsed)
                
                logger.info(f"Successfully parsed document with {len(parsed.sections)} sections")
                return parsed
        
        except Exception as e:
            logger.error(f"Failed to import document {file_path}: {e}")
        
        return None
    
    def _parse_pdf(self, file_path: Path) -> Optional[ParsedDocument]:
        """Parse PDF document"""
        try:
            import PyPDF2
            
            document_id = f"doc_pdf_{file_path.stem.replace(' ', '_')}"
            parsed = ParsedDocument(
                document_id=document_id,
                format=DocumentFormat.PDF,
                title=file_path.stem,
                source_path=str(file_path)
            )
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                # Extract text from all pages
                full_text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text:
                        full_text += f"\n--- Page {page_num + 1} ---\n{text}"
                
                parsed.content = full_text
                
                # Extract metadata
                if pdf_reader.metadata:
                    parsed.metadata = {
                        "author": pdf_reader.metadata.get('/Author', ''),
                        "title": pdf_reader.metadata.get('/Title', ''),
                        "creation_date": pdf_reader.metadata.get('/CreationDate', '')
                    }
                    parsed.author = parsed.metadata.get("author")
            
            # Segment into sections
            self._segment_content(parsed)
            
            return parsed
        
        except ImportError:
            logger.warning("PyPDF2 not installed, skipping PDF parsing")
            return None
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            return None
    
    def _parse_docx(self, file_path: Path) -> Optional[ParsedDocument]:
        """Parse DOCX document"""
        try:
            from docx import Document as DocxDocument
            
            document_id = f"doc_docx_{file_path.stem.replace(' ', '_')}"
            parsed = ParsedDocument(
                document_id=document_id,
                format=DocumentFormat.DOCX,
                title=file_path.stem,
                source_path=str(file_path)
            )
            
            doc = DocxDocument(file_path)
            
            # Extract text and structure
            full_text = ""
            section_id = 0
            
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text += para.text + "\n"
                    
                    # Detect headings
                    style = para.style.name.lower()
                    if 'heading' in style:
                        level = int(style.split()[-1]) if style.split()[-1].isdigit() else 1
                        section = DocumentSection(
                            section_id=f"sec_{section_id}",
                            section_type=SectionType.HEADING if level == 1 else SectionType.SUBHEADING,
                            title=para.text,
                            content="",
                            level=level,
                            line_range=(section_id, section_id)
                        )
                        parsed.sections.append(section)
                        section_id += 1
                    elif 'paragraph' in style or not style:
                        section = DocumentSection(
                            section_id=f"sec_{section_id}",
                            section_type=SectionType.PARAGRAPH,
                            title="",
                            content=para.text,
                            level=0,
                            line_range=(section_id, section_id)
                        )
                        parsed.sections.append(section)
                        section_id += 1
            
            parsed.content = full_text
            
            # Extract author if available
            if doc.core_properties.author:
                parsed.author = doc.core_properties.author
            
            return parsed
        
        except ImportError:
            logger.warning("python-docx not installed, skipping DOCX parsing")
            return None
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            return None
    
    def _parse_markdown(self, file_path: Path) -> Optional[ParsedDocument]:
        """Parse Markdown document"""
        try:
            content = file_path.read_text(encoding='utf-8')
            
            document_id = f"doc_md_{file_path.stem.replace(' ', '_')}"
            parsed = ParsedDocument(
                document_id=document_id,
                format=DocumentFormat.MARKDOWN,
                title=file_path.stem,
                source_path=str(file_path),
                content=content
            )
            
            # Parse markdown structure
            lines = content.split('\n')
            section_id = 0
            current_section = None
            
            for line_num, line in enumerate(lines):
                # Detect headings
                heading_match = re.match(r'^(#+)\s+(.+)$', line)
                if heading_match:
                    level = len(heading_match.group(1))
                    title = heading_match.group(2)
                    
                    section = DocumentSection(
                        section_id=f"sec_{section_id}",
                        section_type=SectionType.HEADING if level <= 2 else SectionType.SUBHEADING,
                        title=title,
                        content="",
                        level=level,
                        line_range=(line_num, line_num)
                    )
                    parsed.sections.append(section)
                    current_section = section_id
                    section_id += 1
                
                elif line.strip() and current_section is not None:
                    # Add content to current section
                    if parsed.sections and section_id - 1 < len(parsed.sections):
                        parsed.sections[section_id - 1].content += line + "\n"
            
            return parsed
        
        except Exception as e:
            logger.error(f"Markdown parsing error: {e}")
            return None
    
    def _parse_html(self, file_path: Path) -> Optional[ParsedDocument]:
        """Parse HTML document"""
        try:
            from html.parser import HTMLParser
            
            content = file_path.read_text(encoding='utf-8')
            
            document_id = f"doc_html_{file_path.stem.replace(' ', '_')}"
            parsed = ParsedDocument(
                document_id=document_id,
                format=DocumentFormat.HTML,
                title=file_path.stem,
                source_path=str(file_path)
            )
            
            # Simple text extraction from HTML
            import re
            text = re.sub('<[^<]+?>', '', content)
            parsed.content = text
            
            return parsed
        
        except Exception as e:
            logger.error(f"HTML parsing error: {e}")
            return None
    
    def _parse_text(self, file_path: Path) -> Optional[ParsedDocument]:
        """Parse plain text document"""
        try:
            content = file_path.read_text(encoding='utf-8')
            
            document_id = f"doc_txt_{file_path.stem.replace(' ', '_')}"
            parsed = ParsedDocument(
                document_id=document_id,
                format=DocumentFormat.TEXT,
                title=file_path.stem,
                source_path=str(file_path),
                content=content
            )
            
            # Simple segmentation by paragraphs
            paragraphs = content.split('\n\n')
            for i, para in enumerate(paragraphs):
                if para.strip():
                    section = DocumentSection(
                        section_id=f"sec_{i}",
                        section_type=SectionType.PARAGRAPH,
                        title="",
                        content=para,
                        level=0
                    )
                    parsed.sections.append(section)
            
            return parsed
        
        except Exception as e:
            logger.error(f"Text parsing error: {e}")
            return None
    
    def _segment_content(self, doc: ParsedDocument):
        """Segment document content into sections"""
        # Simple heuristic: split by blank lines and common separators
        lines = doc.content.split('\n')
        section_id = 0
        current_content = ""
        
        for line in lines:
            if line.startswith('---') or line.startswith('===') or not line.strip():
                if current_content.strip():
                    section = DocumentSection(
                        section_id=f"sec_{section_id}",
                        section_type=SectionType.PARAGRAPH,
                        title="",
                        content=current_content,
                        level=0
                    )
                    doc.sections.append(section)
                    section_id += 1
                    current_content = ""
            else:
                current_content += line + "\n"
        
        if current_content.strip():
            section = DocumentSection(
                section_id=f"sec_{section_id}",
                section_type=SectionType.PARAGRAPH,
                title="",
                content=current_content,
                level=0
            )
            doc.sections.append(section)
    
    def _extract_entities(self, doc: ParsedDocument):
        """Extract entities and keywords from document"""
        all_text = doc.content.lower()
        
        # Extract using patterns
        for entity_type, pattern in self.ENTITY_PATTERNS.items():
            matches = re.findall(pattern, doc.content)
            if matches:
                if isinstance(matches[0], tuple):
                    # Handle multiple capture groups
                    matches = [m for group in matches for m in group if m]
                doc.entities[entity_type] = list(set(matches))
        
        # Simple keyword extraction (frequent words)
        words = re.findall(r'\b[a-z]{4,}\b', all_text)
        from collections import Counter
        common_words = {"the", "and", "that", "with", "for", "this", "from", "are", "but", "not"}
        word_counts = Counter(w for w in words if w not in common_words)
        doc.keywords = [w for w, _ in word_counts.most_common(20)]
    
    def _save_to_vault(self, doc: ParsedDocument):
        """Save parsed document to VaultCore"""
        try:
            from vault_core import VaultEntry, VaultEntryType
            
            # Create main document entry
            entry = VaultEntry(
                entry_id=doc.document_id,
                entry_type=VaultEntryType.DOCUMENTATION,
                title=doc.title,
                content=doc.content[:5000],  # Limit content size
                summary=f"Imported {doc.format.value} document with {len(doc.sections)} sections",
                tags=["graphify", doc.format.value] + doc.keywords[:5],
                created_by="GraphifyEngine"
            )
            
            self.vault_core.store(entry)
            
            # Create section entries
            for section in doc.sections:
                section_entry = VaultEntry(
                    entry_id=f"{doc.document_id}_{section.section_id}",
                    entry_type=VaultEntryType.DOCUMENTATION,
                    title=section.title or f"Section {section.section_id}",
                    content=section.content,
                    summary=f"Section from {doc.title}",
                    tags=["section", doc.format.value] + section.keywords,
                    created_by="GraphifyEngine"
                )
                
                self.vault_core.store(section_entry)
            
            logger.info(f"Saved document {doc.document_id} to VaultCore")
        
        except Exception as e:
            logger.error(f"Failed to save document to VaultCore: {e}")
    
    def batch_import(self, directory: Path, pattern: str = "*.md") -> List[ParsedDocument]:
        """Import multiple documents from directory"""
        documents = []
        
        for file_path in directory.glob(pattern):
            if file_path.is_file():
                doc = self.import_document(file_path)
                if doc:
                    documents.append(doc)
        
        logger.info(f"Batch imported {len(documents)} documents")
        return documents
    
    def export_parsed(self, output_path: Path) -> bool:
        """Export all parsed documents to JSON"""
        try:
            data = {
                doc_id: doc.to_dict() 
                for doc_id, doc in self.parsed_documents.items()
            }
            
            output_path.write_text(
                json.dumps(data, indent=2, default=str),
                encoding='utf-8'
            )
            logger.info(f"Exported {len(data)} documents to {output_path}")
            return True
        except Exception as e:
            logger.error(f"Export failed: {e}")
            return False


if __name__ == "__main__":
    # Demo
    logging.basicConfig(level=logging.DEBUG)
    
    engine = GraphifyEngine()
    
    # Try to import a sample document
    test_file = Path("README.md")
    if test_file.exists():
        doc = engine.import_document(test_file)
        if doc:
            print(f"\nParsed: {doc.title}")
            print(f"Sections: {len(doc.sections)}")
            print(f"Keywords: {doc.keywords[:5]}")
